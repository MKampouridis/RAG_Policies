#!/usr/bin/env python3
"""CLI: ingest a LOCAL document (.docx) that the crawler cannot reach.

Why this exists separately from run_ingest.py: that pipeline is URL-driven
end to end - it crawls essex.ac.uk, and every document it indexes is a public
PDF or an explicitly-exempted HTML page. Some documents that answer real
questions are never published there: staff-facing handbooks circulated
internally, for instance. They have no URL to crawl and no PDF to fetch.

This script reuses the SAME downstream pipeline (classify -> upsert_document
-> manifest -> text cache), so a locally-ingested document is indistinguishable
from a crawled one once stored. Only the acquisition step differs.

Two design points worth knowing:

  * The document is COPIED into data/local_documents/ and keyed by the URL
    path it is served from (/documents/<file>), not by its path on disk. The
    UI's source modal sets a link href to the stored URL (static/app.js
    paintModal), so a key like "file:///Users/..." would render a citation
    that points at nothing - the exact defect the "Repair inline citations"
    commit fixed. src/app.py mounts data/local_documents at /documents so the
    citation resolves for anyone using the app.

  * run_ingest.py does NOT prune documents it did not crawl (delete_document
    is only called for a crawled item that no longer qualifies), so a
    locally-ingested document survives future crawls. It is NOT re-fetched or
    re-checked by them either: re-run this script when the source file changes.

Usage:
    python ingest_local.py <path.docx> [--title "..."] [--doc-type policy]
                           [--department X] [--year 2025-26] [--internal]

    --internal marks the document as not publicly published, so answers and
    audits can tell it apart from crawled public policy.
"""

import argparse
import hashlib
import shutil
import sys
import time
from pathlib import Path

import docx

from reembed import compute_current_flags
from run_ingest import (
    TEXT_CACHE_DIR,
    _refuse_if_server_live,
    _sync_family_siblings,
    load_manifest,
    save_manifest,
)
from src.ingest import upsert_document, url_hash
from src.relevance import classify

LOCAL_DOCS_DIR = Path("data/local_documents")
URL_PREFIX = "/documents/"


def _iter_block_text(document) -> list[str]:
    """Paragraphs and tables in document order.

    python-docx exposes .paragraphs and .tables as separate flat lists, which
    loses their interleaving - and this document's substance is largely in
    appendix tables sitting between explanatory paragraphs. Walking the body
    XML keeps reading order, so a chunk boundary can't land between a table
    and the sentence that introduces it.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    out = []
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            out.append(Paragraph(child, document).text)
        elif tag == "tbl":
            table = Table(child, document)
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                # de-duplicate horizontally merged cells, which python-docx
                # reports once per underlying grid column
                deduped = [c for i, c in enumerate(cells) if i == 0 or c != cells[i - 1]]
                line = " | ".join(c for c in deduped if c)
                if line:
                    out.append(line)
    return out


def extract_docx_text(path: Path) -> str:
    """Readable text from a .docx, minus Word's table-of-contents plumbing.

    TOC entries survive as literal field text ("Introduction PAGEREF
    _Toc234314569 \\h 2"). They are pure noise for retrieval - a list of
    headings the body already contains, carrying page numbers that mean
    nothing once chunked - and they would otherwise be the document's most
    heading-dense chunk, which is exactly the shape that wins on identity
    queries while answering nothing.
    """
    document = docx.Document(str(path))
    lines = []
    for raw in _iter_block_text(document):
        line = raw.strip()
        if not line:
            continue
        if "PAGEREF" in line or line.startswith("TOC \\"):
            continue
        lines.append(line)
    return "\n".join(lines)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("path", help="path to the .docx to ingest")
    ap.add_argument("--title", help="document title (default: derived from filename)")
    ap.add_argument("--doc-type", help="override the classifier's doc_type")
    ap.add_argument("--department", help="override the classifier's department")
    ap.add_argument("--year", help="override the classifier's academic_year, e.g. 2025-26")
    ap.add_argument("--internal", action="store_true",
                    help="mark as an internal (unpublished) document")
    ap.add_argument("--name", help="filename to serve it as (default: the source filename)")
    args = ap.parse_args()

    src = Path(args.path).expanduser()
    if not src.exists():
        print(f"no such file: {src}", file=sys.stderr)
        return 1
    if src.suffix.lower() != ".docx":
        print(f"only .docx is supported here; {src.suffix} files that are public "
              f"should go through run_ingest.py by URL", file=sys.stderr)
        return 1

    # Writing to Chroma while the server holds it open leaves that process
    # serving a stale in-memory index - the same reason run_ingest refuses.
    _refuse_if_server_live()

    text = extract_docx_text(src)
    if not text.strip():
        print("no extractable text", file=sys.stderr)
        return 1

    served_name = args.name or src.name
    url = URL_PREFIX + served_name
    title = args.title or src.stem.replace("_", " ").replace("-", " ")

    LOCAL_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    dest = LOCAL_DOCS_DIR / served_name
    shutil.copy2(src, dest)

    TEXT_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_path = TEXT_CACHE_DIR / f"{url_hash(url)}.txt"
    cache_path.write_text(text, encoding="utf-8")

    content_hash = hashlib.sha256(dest.read_bytes()).hexdigest()

    print(f"extracted {len(text):,} chars from {src.name}")
    print(f"  served as: {url}")

    manifest = load_manifest()
    documents = manifest["documents"]
    prior = documents.get(url)
    if prior and prior.get("content_hash") == content_hash:
        print("  unchanged since last ingest - nothing to do")
        return 0

    try:
        decision = classify(title, url, text)
    except Exception as exc:
        print(f"  classification failed ({exc}); falling back to explicit metadata")
        decision = {"keep": True, "doc_type": args.doc_type or "policy",
                    "department": args.department, "academic_year": args.year,
                    "reason": "classification failed; kept on explicit request"}

    # An operator ingesting a file by hand has named the document deliberately;
    # the classifier is a convenience here, not the authority it is on a crawl
    # of thousands of pages.
    if args.doc_type:
        decision["doc_type"] = args.doc_type
        decision["keep"] = True
    if args.department:
        decision["department"] = args.department
    if args.year:
        decision["academic_year"] = args.year

    print(f"  classified: keep={decision['keep']} type={decision.get('doc_type')} "
          f"dept={decision.get('department')} year={decision.get('academic_year')}")

    changed_at = time.time() if prior else None
    entry = {
        "url": url,
        "title": title,
        "content_type": "docx",
        "content_hash": content_hash,
        "text_cache_path": str(cache_path),
        "source": "local",
        "local_path": str(dest),
        **({"published": False} if args.internal else {}),
        **({"content_changed_at": changed_at} if changed_at else {}),
        **decision,
    }
    documents[url] = entry

    if not decision["keep"]:
        save_manifest(manifest)
        print(f"  rejected ({decision.get('reason', '')}) - not indexed")
        return 0

    flags = compute_current_flags(documents)
    metadata = {
        "title": title,
        "doc_type": decision["doc_type"],
        "department": decision.get("department"),
        "academic_year": decision.get("academic_year"),
        "is_current": flags[url],
    }
    n_chunks = upsert_document(url, text, metadata)
    entry["chunk_count"] = n_chunks
    if flags[url]:
        _sync_family_siblings(url, flags)
    save_manifest(manifest)

    print(f"  INDEXED: {n_chunks} chunks, is_current={flags[url]}")
    print("\nNow run the post-ingest checks:")
    print("  python eval/stale_index_audit.py")
    print("  python eval/check_benchmark_stamp.py")
    print("  python audit_family_aliases.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())
